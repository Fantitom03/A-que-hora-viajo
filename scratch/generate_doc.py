import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_document():
    doc = docx.Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Estilos globales
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(33, 37, 41) # Charcoal
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(8)

    # Colores principales
    PRIMARY_COLOR = RGBColor(16, 44, 87)    # Deep Blue
    SECONDARY_COLOR = RGBColor(74, 85, 104) # Slate Gray
    
    # ------------------ PORTADA / TÍTULO ------------------
    p_pre = doc.add_paragraph()
    p_pre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pre = p_pre.add_run("TRABAJO PRÁCTICO ACADÉMICO")
    run_pre.bold = True
    run_pre.font.size = Pt(10)
    run_pre.font.color.rgb = SECONDARY_COLOR
    p_pre.paragraph_format.space_after = Pt(2)
    p_pre.paragraph_format.space_before = Pt(40)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Caso de Estudio: Sistema de Gestión y Monitoreo de Terminal de Ómnibus\n\"A Qué Hora Viajo\"")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = PRIMARY_COLOR
    title_p.paragraph_format.space_after = Pt(6)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("Documento de Alcance y Especificación Funcional")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = SECONDARY_COLOR
    subtitle_p.paragraph_format.space_after = Pt(150)

    # Datos del Proyecto/Autores
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_univ = p_info.add_run("Universidad Nacional de Catamarca\n")
    run_univ.bold = True
    run_univ.font.size = Pt(11)
    
    run_fac = p_info.add_run("Facultad de Tecnología y Ciencias Aplicadas\nIngeniería Informática - Cátedra: Desarrollo de APIs\n\n")
    run_fac.font.size = Pt(10)
    run_fac.font.color.rgb = SECONDARY_COLOR
    
    run_auth = p_info.add_run("Autores:\n- Centeno Joaquín\n- Leiva Francisco\n\n")
    run_auth.bold = True
    run_auth.font.size = Pt(11)
    
    run_date = p_info.add_run("Fecha: Junio 2026")
    run_date.font.size = Pt(10)
    run_date.italic = True
    
    doc.add_page_break()

    # ------------------ CONTENIDO ------------------
    
    # 1. Introducción
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Introducción y Contexto")
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1_run.font.name = 'Arial'
    h1_run.bold = True
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph(
        "El presente caso de estudio detalla la arquitectura, el alcance y las especificaciones funcionales de la "
        "plataforma backend y frontend \"A Qué Hora Viajo\". Este sistema está diseñado para abordar una problemática "
        "común en terminales de colectivos de media y larga distancia: la falta de previsibilidad e información oportuna "
        "hacia los pasajeros respecto a los horarios de partida, arribos programados, demoras imprevistas y condiciones "
        "climáticas en las paradas de destino."
    )
    
    doc.add_paragraph(
        "Tradicionalmente, las terminales de transporte muestran horarios fijos en pantallas físicas o carteleras. "
        "Sin embargo, factores como el tráfico, roturas mecánicas o inclemencias climáticas alteran de manera constante "
        "los tiempos de viaje. \"A Qué Hora Viajo\" centraliza la información operativa de la terminal, permitiendo a las "
        "empresas de transporte gestionar y reportar estados en tiempo real, mientras que proporciona a los pasajeros "
        "herramientas avanzadas de búsqueda y previsión."
    )

    # 2. Objetivos del Sistema
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Objetivos del Sistema")
    h2_run.font.color.rgb = PRIMARY_COLOR
    h2_run.font.name = 'Arial'
    h2_run.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "El sistema persigue los siguientes objetivos generales:"
    )
    
    bullets = [
        "Proporcionar una API REST robusta y segura para que las empresas y administradores gestionen terminales, colectivos y horarios.",
        "Ofrecer a los pasajeros un buscador intuitivo que permita consultar recorridos y obtener de manera dinámica el horario estimado de llegada a su parada de interés.",
        "Automatizar e integrar servicios de terceros como OpenWeatherMap para predecir el clima en el destino exacto al momento del arribo del colectivo.",
        "Optimizar la carga y validación de ubicaciones mediante el uso de mapas interactivos y servicios de geolocalización inversa como OpenStreetMap Nominatim.",
        "Implementar un módulo de pantalla de terminal en tiempo real que liste las salidas recientes y los arribos programados del día actual."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # 3. Roles de Usuario y Permisos
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Roles de Usuario y Gestión de Acceso")
    h3_run.font.color.rgb = PRIMARY_COLOR
    h3_run.font.name = 'Arial'
    h3_run.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Para garantizar la integridad de los datos y restringir acciones sensibles, se definen cuatro roles principales con un control de acceso basado en tokens JSON Web Tokens (JWT):"
    )

    roles = [
        ("Administrador del Sistema (Superusuario):", "Posee acceso total e irrestricto a todos los recursos de la API. Sus tareas principales incluyen dar de alta y baja terminales físicas, registrar nuevas empresas de transporte y asignar encargados generales para cada empresa."),
        ("Encargado de Empresa:", "Es un empleado con jerarquía administrativa dentro de una empresa específica. Puede registrar nuevos empleados para su empresa (boleteros o personal de ventanilla), actualizar la información corporativa y gestionar la flota de viajes programados."),
        ("Empleado de Ventanilla:", "Personal operativo asignado a una empresa de transporte. Sus funciones se centran en registrar viajes diarios, definir plataformas de embarque y reportar estados operativos en tiempo real (ej. demoras de colectivos, cancelaciones, cambios de plataforma)."),
        ("Pasajeros / Público General:", "Usuarios finales que consultan el sistema. No requieren autenticación para buscar viajes o visualizar la pantalla de arribos. Sin embargo, si se registran, pueden configurar perfiles, asociar números de contacto y habilitar notificaciones/alertas de viaje.")
    ]
    for r_title, r_desc in roles:
        p = doc.add_paragraph(style='List Bullet')
        run_title = p.add_run(r_title + " ")
        run_title.bold = True
        p.add_run(r_desc)

    # 4. Módulos y Funcionalidades Clave
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Alcance Funcional e Integraciones")
    h4_run.font.color.rgb = PRIMARY_COLOR
    h4_run.font.name = 'Arial'
    h4_run.bold = True
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "El sistema destaca por la implementación de las siguientes soluciones tecnológicas y lógicas de negocio:"
    )

    func = [
        ("Validación y Sanitización Geográfica (OpenStreetMap Nominatim):", "Al crear o modificar una Ubicación geográfica, el backend intercepta el nombre ingresado (ej. 'Terminal Mendoza') y realiza una petición HTTP al servidor de Nominatim. El sistema valida su existencia, corrige errores ortográficos y obtiene la latitud y longitud oficiales con una precisión de 6 decimales. Esto previene registros erróneos y asegura coordenadas exactas para consultas posteriores."),
        ("Búsqueda Inteligente de Recorridos:", "El buscador de pasajes procesa búsquedas flexibles y tolerantes a acentos (accent-insensitive) e inclusive parciales (ej. buscar 'tucuman' devuelve 'San Miguel de Tucumán, Tucumán, Argentina'). Filtra los viajes por día de la semana operativo y comprueba que el destino coincida con alguna de las paradas intermedias o finales del viaje."),
        ("Cálculo Dinámico de Horarios y Demoras:", "El horario de arribo a una parada no es estático. El sistema calcula la hora de llegada sumando al horario de salida del viaje el tiempo estimado de viaje transcurrido (DurationField) para esa parada específica, sumándole de forma dinámica cualquier demora (tiempo_demora) registrada por la empresa para el viaje en la fecha de consulta."),
        ("Predicción Climática Inteligente (OpenWeatherMap):", "Una vez calculada la hora exacta de arribo estimada del pasajero a su parada de destino, el backend calcula la fecha correspondiente y realiza una petición a OpenWeather API. Si el viaje ocurre dentro de los próximos 5 días, obtiene el pronóstico horario exacto del destino. Si es a mayor plazo, obtiene las condiciones generales del clima actual de la zona. De esta manera, el usuario sabe con qué condiciones meteorológicas se encontrará al bajar del colectivo."),
        ("Evaluación de Estados 'Lazy' en Tiempo Real:", "Para evitar que el personal de la empresa deba marcar de forma manual la finalización de cada viaje, el sistema realiza una evaluación perezosa (lazy evaluation) del estado. Si la hora actual del servidor supera a la hora de arribo estimada final del viaje más un margen de tolerancia (1 hora), el backend actualiza de forma automática el estado a 'FINALIZADO'. Si el horario de embarque y demora ya transcurrieron y el estado era 'A_TIEMPO', cambia automáticamente a 'EN_VIAJE'.")
    ]
    for f_title, f_desc in func:
        p = doc.add_paragraph()
        run_title = p.add_run(f_title + "\n")
        run_title.bold = True
        run_title.font.color.rgb = SECONDARY_COLOR
        p.add_run(f_desc)
        p.paragraph_format.space_after = Pt(12)

    # 5. Modelo de Datos y Entidades principales
    h5 = doc.add_heading(level=1)
    h5_run = h5.add_run("5. Modelo de Datos y Arquitectura de la API")
    h5_run.font.color.rgb = PRIMARY_COLOR
    h5_run.font.name = 'Arial'
    h5_run.bold = True
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "El esquema relacional de la base de datos se modela bajo las siguientes entidades principales, utilizando PostgreSQL como motor y UUIDs autogenerados para garantizar claves primarias seguras y no correlativas:"
    )

    models_list = [
        ("Terminal:", "Representa el nodo físico. Almacena el nombre de la terminal y la cantidad de plataformas disponibles para el arribo de colectivos."),
        ("Empresa:", "Empresas de transporte registradas en la terminal. Contiene nombre, ventanilla física asignada y está vinculada a una Terminal determinada."),
        ("Usuario / Empleado / Pasajero:", "Gestión extendida del modelo de usuarios de Django. El DNI actúa como campo de validación único. Los empleados se asocian a una empresa y poseen roles de ENCARGADO o VENTANILLA. Los pasajeros almacenan preferencias y teléfonos de contacto."),
        ("Ubicacion:", "Puntos geográficos oficiales normalizados mediante OSM Nominatim con latitud y longitud decimal."),
        ("Viaje:", "Esquema troncal de frecuencia y horarios de colectivos. Define la empresa operadora, la hora de embarcación de origen, la duración total, y optimiza los días operativos y plataformas asignadas utilizando el tipo de dato ArrayField de PostgreSQL."),
        ("Parada:", "Entidad intermedia que relaciona un viaje con sus múltiples ubicaciones a lo largo del recorrido. Registra el orden de la parada, el tiempo relativo transcurrido desde la salida y el precio del pasaje hasta ese punto."),
        ("EstadoViajeDiario:", "Mapea las novedades operativas del día. Mantiene el estado en tiempo real (A_TIEMPO, EN_VIAJE, ABORDANDO, DEMORADO, FINALIZADO, CANCELADO) para una fecha dada, la duración de la demora acumulada y el motivo que la originó.")
    ]
    for m_title, m_desc in models_list:
        p = doc.add_paragraph(style='List Bullet')
        run_title = p.add_run(m_title + " ")
        run_title.bold = True
        p.add_run(m_desc)

    doc.add_paragraph()

    # 6. Tabla de Endpoints de la API
    h6 = doc.add_heading(level=1)
    h6_run = h6.add_run("6. Especificación de Endpoints del Backend")
    h6_run.font.color.rgb = PRIMARY_COLOR
    h6_run.font.name = 'Arial'
    h6_run.bold = True
    h6.paragraph_format.space_before = Pt(12)
    h6.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "A continuación se detallan los principales endpoints REST expuestos por el backend para el consumo del frontend y sistemas externos:"
    )

    # Crear tabla de Endpoints
    endpoints_data = [
        ("Método", "Endpoint", "Descripción / Parámetros", "Acceso"),
        ("POST", "/api/token/", "Obtiene token de acceso JWT y refresh token.", "Público"),
        ("POST", "/api/token/refresh/", "Renueva el token de acceso JWT.", "Público"),
        ("GET", "/api/terminal/", "Lista las terminales físicas configuradas.", "Autenticado"),
        ("POST", "/api/terminal/", "Da de alta una nueva terminal.", "Superuser"),
        ("POST", "/api/empleado/", "Registra un empleado asociado a una empresa.", "Encargado/Superuser"),
        ("POST", "/api/pasajero/", "Registra de forma pública un nuevo pasajero.", "Público"),
        ("POST", "/api/ubicacion/", "Registra ubicación georeferenciada. Llama a Nominatim.", "Empleado/Superuser"),
        ("GET", "/api/viaje/", "Lista viajes programados. Filtra por empresa según rol.", "Público / Filtrado"),
        ("POST", "/api/viaje/", "Crea un nuevo viaje vinculándolo a la empresa del empleado.", "Personal Empresa"),
        ("GET", "/api/viaje/buscar_viajes/", "Buscador de viajes. Parámetros: 'destino', 'dia', 'terminal_id'. Integra clima y demoras.", "Público"),
        ("GET", "/api/viaje/pantalla_terminal/", "Viajes programados y arribos de hoy. Parámetro: 'terminal_id'.", "Público"),
        ("POST", "/api/viaje/{id}/actualizar_estado_diario/", "Registra demora, motivo o cambia estado del viaje para una fecha específica.", "Personal Empresa")
    ]

    table = doc.add_table(rows=len(endpoints_data), cols=4)
    table.style = 'Table Grid'
    
    # Configurar anchos de columna
    col_widths = [Inches(1.0), Inches(2.2), Inches(2.3), Inches(1.5)]
    
    for r_idx, row in enumerate(table.rows):
        # Configurar alto de fila para cabecera
        if r_idx == 0:
            trPr = row._tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '360')
            trPr.append(trHeight)
            
        for c_idx, cell in enumerate(row.cells):
            cell.width = col_widths[c_idx]
            cell.text = endpoints_data[r_idx][c_idx]
            
            # Estilos de cabecera
            if r_idx == 0:
                set_cell_background(cell, "102C57") # Deep Blue hex
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(10)
            else:
                # Estilo filas de datos
                p = cell.paragraphs[0]
                p.runs[0].font.size = Pt(9.5)
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    method = endpoints_data[r_idx][0]
                    # Resaltar métodos HTTP
                    if method == "GET":
                        p.runs[0].font.color.rgb = RGBColor(40, 167, 69) # Green
                        p.runs[0].font.bold = True
                    elif method == "POST":
                        p.runs[0].font.color.rgb = RGBColor(0, 123, 255) # Blue
                        p.runs[0].font.bold = True
                elif c_idx == 1:
                    p.runs[0].font.name = 'Courier New'

    # Conclusión
    doc.add_paragraph()
    h7 = doc.add_heading(level=1)
    h7_run = h7.add_run("7. Conclusiones y Trabajo Futuro")
    h7_run.font.color.rgb = PRIMARY_COLOR
    h7_run.font.name = 'Arial'
    h7_run.bold = True
    h7.paragraph_format.space_before = Pt(12)
    h7.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "El sistema \"A Qué Hora Viajo\" demuestra cómo las integraciones modernas y el diseño de APIs inteligentes "
        "mejoran drásticamente la calidad y la previsibilidad de los servicios públicos de transporte. La automatización "
        "del cálculo meteorológico y geográfico alivia la carga de los operarios, garantizando datos limpios y útiles."
    )
    doc.add_paragraph(
        "Como líneas de trabajo futuro se plantea la incorporación de notificaciones automáticas vía correo electrónico/SMS "
        "para avisar a los pasajeros si su colectivo se demora más de un umbral establecido, así como la incorporación de "
        "seguimiento GPS real de las unidades a través de dispositivos de a bordo."
    )

    # Guardar
    output_path = "Alcance_Caso_Estudio_Terminal_Omnibus.docx"
    doc.save(output_path)
    print(f"Documento creado exitosamente en: {output_path}")

if __name__ == "__main__":
    create_document()
